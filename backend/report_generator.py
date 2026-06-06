from docx import Document
from docx.shared import Inches
from pathlib import Path
import uuid
from datetime import datetime


class ReportGenerator:
    def generate_report(self, filename, question, answer, table, chart_path=None):
        reports_dir = Path(__file__).resolve().parent.parent / "data" / "reports"
        reports_dir.mkdir(parents=True, exist_ok=True)

        report_name = f"business_report_{uuid.uuid4()}.docx"
        report_path = reports_dir / report_name

        doc = Document()

        doc.add_heading("AI Business Analytics Report", level=1)

        doc.add_paragraph(f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Source File: {filename}")
        doc.add_paragraph(f"Business Question: {question}")

        doc.add_heading("AI Analysis", level=2)
        doc.add_paragraph(answer)

        if table:
            doc.add_heading("Forecast Results" if "forecast" in question.lower() else "Calculated Results", level=2)

            headers = list(table[0].keys())

            doc_table = doc.add_table(rows=1, cols=len(headers))
            doc_table.style = "Table Grid"

            for i, header in enumerate(headers):
                doc_table.rows[0].cells[i].text = header

            for row in table:
                cells = doc_table.add_row().cells

                for i, header in enumerate(headers):
                    cells[i].text = str(row[header])

        if chart_path:
            doc.add_heading("Forecast Chart" if "forecast" in question.lower() else "Analytics Chart", level=2)

            base_dir = Path(__file__).resolve().parent.parent

            if chart_path.startswith("/charts/"):
                image_path = (
                    base_dir
                    / "data"
                    / "charts"
                    / Path(chart_path).name
                )
            else:
                image_path = Path(chart_path)

            if image_path.exists():
                doc.add_picture(
                    str(image_path),
                    width=Inches(5.8)
                )

        doc.add_heading("Executive Summary", level=2)

        if "forecast" in question.lower():
            doc.add_paragraph(
                "This report contains AI-generated revenue forecasts "
                "for future business planning and decision making."
            )
        else:
            doc.add_paragraph(
                "This report contains analytical insights generated "
                "from uploaded business data."
            )

        return f"/reports/{report_name}"